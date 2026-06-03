from pathlib import Path
from docx import Document


template = Path(r"D:\__easyHelper__\学习资料\领域软件工程\领域软件工程-报告模板-2026.docx")
doc = Document(template)

print("SECTIONS", len(doc.sections))
print("PARAGRAPHS", len(doc.paragraphs))
for index, paragraph in enumerate(doc.paragraphs):
    text = paragraph.text.strip()
    if text:
        print(f"P{index:03d}\t{paragraph.style.name}\t{text}")

print("TABLES", len(doc.tables))
for table_index, table in enumerate(doc.tables):
    print(f"TABLE {table_index} rows={len(table.rows)} cols={len(table.columns)}")
    for row_index, row in enumerate(table.rows):
        values = [cell.text.replace("\n", " / ").strip() for cell in row.cells]
        print(f"R{row_index:02d}\t" + "\t".join(values))

print("INLINE_SHAPES", len(doc.inline_shapes))
