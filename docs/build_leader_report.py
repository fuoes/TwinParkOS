from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = Path(r"D:\__easyHelper__\学习资料\领域软件工程\领域软件工程-报告模板-2026.docx")
OUTPUT = ROOT / "docs" / "智慧产业园区数字孪生可视化管控平台-组长个人报告.docx"
ASSETS = ROOT / "docs" / "report-assets"

PROJECT_NAME = "智慧产业园区数字孪生可视化管控平台"
PLACEHOLDER_CLASS = "[请填写年级专业/班级]"
PLACEHOLDER_NAME = "[请填写组长姓名]"
PLACEHOLDER_MEMBER_2 = "[请填写成员二姓名]"
PLACEHOLDER_MEMBER_3 = "[请填写成员三姓名]"

FONT_BODY = "宋体"
FONT_HEADING = "黑体"
BODY_SIZE = Pt(12)
TABLE_SIZE = Pt(10.5)


def set_run_font(run, name=FONT_BODY, size=BODY_SIZE, bold=None, color=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = size
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def replace_paragraph(paragraph, text, bold_prefix=None):
    clear_paragraph(paragraph)
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.add_paragraph()
    new_para._p.getparent().remove(new_para._p)
    new_p.addnext(new_para._p)
    paragraph._p.getparent().remove(new_p)
    if style:
        new_para.style = style
    if text:
        run = new_para.add_run(text)
        set_run_font(run)
    return new_para


def remove_from_paragraph(doc, paragraph):
    body = doc._element.body
    start = paragraph._p
    deleting = False
    for child in list(body):
        if child is start:
            deleting = True
        if deleting and child.tag != qn("w:sectPr"):
            body.remove(child)


def set_body_format(paragraph, first_line=True, space_after=Pt(6), line_spacing=1.5):
    fmt = paragraph.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.line_spacing = line_spacing
    fmt.space_after = space_after
    fmt.space_before = Pt(0)
    fmt.first_line_indent = Pt(24) if first_line else Pt(0)
    fmt.widow_control = True


def add_body(doc, text, first_line=True, bold_prefix=None, space_after=Pt(6)):
    paragraph = doc.add_paragraph()
    set_body_format(paragraph, first_line=first_line, space_after=space_after)
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = paragraph.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_heading(doc, text, level):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, FONT_HEADING, Pt(16 if level == 1 else 14), bold=True)
    return paragraph


def add_caption(doc, text, kind="figure"):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_with_next = False
    run = paragraph.add_run(text)
    set_run_font(run, FONT_BODY, Pt(10.5))
    return paragraph


def add_figure(doc, path, caption, width=Inches(6.15)):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=width)
    add_caption(doc, caption)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(str(text))
    set_run_font(run, FONT_BODY, TABLE_SIZE, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, top=True, header_bottom=True, bottom=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        enabled = (edge == "top" and top) or (edge == "bottom" and bottom)
        node.set(qn("w:val"), "single" if enabled else "nil")
        node.set(qn("w:sz"), "12" if enabled else "0")
        node.set(qn("w:color"), "000000")
    if header_bottom and table.rows:
        for cell in table.rows[0].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in("w:tcBorders")
            if tc_borders is None:
                tc_borders = OxmlElement("w:tcBorders")
                tc_pr.append(tc_borders)
            bottom_node = tc_borders.find(qn("w:bottom"))
            if bottom_node is None:
                bottom_node = OxmlElement("w:bottom")
                tc_borders.append(bottom_node)
            bottom_node.set(qn("w:val"), "single")
            bottom_node.set(qn("w:sz"), "8")
            bottom_node.set(qn("w:color"), "000000")


def set_table_widths(table, widths_inches):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_twips = int(sum(widths_inches) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            twips = int(widths_inches[index] * 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(twips))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def mark_header_repeat(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc, caption, headers, rows, widths, aligns=None):
    add_caption(doc, caption, kind="table")
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    set_table_borders(table)
    mark_header_repeat(table.rows[0])
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            align = aligns[index] if aligns else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[index], value, align=align)
    return table


def fill_cover(doc):
    replace_paragraph(doc.paragraphs[17], f"年级专业：   {PLACEHOLDER_CLASS}")
    replace_paragraph(doc.paragraphs[18], f"考生姓名：   {PLACEHOLDER_NAME}")
    table = doc.tables[0]
    set_cell_text(table.cell(0, 1), PROJECT_NAME, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(1, 1), PLACEHOLDER_CLASS, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(2, 1), f"组长：{PLACEHOLDER_NAME}；成员二：{PLACEHOLDER_MEMBER_2}", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(table.cell(3, 1), f"成员三：{PLACEHOLDER_MEMBER_3}", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.core_properties.title = f"{PROJECT_NAME} - 组长个人报告"
    doc.core_properties.subject = "领域软件工程课程报告"
    doc.core_properties.author = PLACEHOLDER_NAME


def build_report():
    copy2(TEMPLATE, OUTPUT)
    doc = Document(OUTPUT)
    fill_cover(doc)

    start = next(p for p in doc.paragraphs if p.text.strip() == "1 项目背景与目标")
    remove_from_paragraph(doc, start)

    add_heading(doc, "1 项目背景与目标", 1)
    add_heading(doc, "1.1 项目背景", 2)
    add_body(doc, "项目应用领域：本项目属于智慧产业园区运营管理与数字孪生可视化领域，面向经开区科创产业园的设备运维、能耗管理、安防管控、环境监测、企业服务和空间资产管理。", bold_prefix="项目应用领域：")
    add_body(doc, "发起方：科创产业园运营管理有限公司。该单位负责园区整体运营、物业运维、企业服务、设施管理和安防管控，需要一个能够统一汇聚空间、设备和事件数据的运营平台。", bold_prefix="发起方：")
    add_body(doc, "驱动因素：随着园区数字化转型持续推进，空间计算和数字孪生技术逐渐成为连接物理空间、物联网数据与业务流程的重要入口。空间计算能够将位置、对象和实时数据关联起来，为园区态势感知与跨系统协同提供基础[1]。", bold_prefix="驱动因素：")
    add_body(doc, "业务痛点：园区内电梯、配电柜、空调、水泵、摄像头、门禁、道闸和环境传感器分散在不同区域，设备状态依赖人工巡检或独立系统查看；水电能耗数据缺少楼栋、企业和设备维度的统一分析；消防、门禁和设备告警难以快速定位，且从告警确认到派单、处理、验收和归档之间缺少闭环；管理层也难以在一个页面内掌握企业入驻、空间出租、设备健康、告警和工单效率。", bold_prefix="业务痛点：")
    add_body(doc, "现有方案及挑战：传统园区系统通常按设备厂家或业务专业分别建设，能够完成局部监控，但跨系统数据关联能力较弱。真实设备接入还会受到协议不统一、接口开放程度、现场网络和数据质量影响。本课程项目没有本地设备接入条件，因此采用“虚拟设备仿真 + 完整业务功能”的实现策略：设备、告警和实时数据由仿真器生成，登录鉴权、权限控制、告警处置、工单闭环、企业服务和系统管理等功能均保持可操作、可测试和可追溯。")

    add_heading(doc, "1.2 项目目标", 2)
    add_body(doc, "本项目建设目标是形成一个以三维数字孪生园区为统一空间入口，以设备、能耗、安防、环境、企业、空间和工单数据为业务支撑的可视化运营管控平台，实现园区运行状态“看得见、管得住、可预警、能联动、可分析、可追溯”。")
    add_body(doc, "系统在虚拟演示模式下需要满足三个层面的目标。第一，提供可视化能力，通过首页驾驶舱和三维一张图展示园区楼栋、设备、告警和关键运营指标；第二，提供业务闭环能力，使告警能够被确认、定位、生成工单、派发、处理、验收并关闭；第三，提供治理能力，通过 JWT 登录、RBAC 权限、系统日志、数据字典和接口管理保证不同角色在授权范围内使用系统。")
    add_body(doc, "作为组长，本人承担 40% 的工作量，主要负责系统总体架构、登录鉴权、安防管控、告警中心、物业工单、企业服务、系统管理、系统集成、自动化测试和 GitHub 版本管理。本人负责的核心价值是打通“安防事件或设备异常 - 告警中心 - 物业工单 - 处置归档”的运营闭环。")

    add_heading(doc, "2 需求分析与系统设计", 1)
    add_heading(doc, "2.1 系统功能模块", 2)
    add_body(doc, "系统包含首页驾驶舱、数字孪生一张图、设备运维、能耗管理、安防管控、环境监测、物业工单、企业服务、空间资产、告警中心、报表中心和系统管理十二个一级业务模块。园区运行数据通过虚拟设备和仿真场景产生，异常事件进入统一告警中心，再根据告警类型进行三维定位、视频联动、人员确认和工单处置，整体流程如图1所示。")
    add_figure(doc, ASSETS / "diagram-1-business-flow.png", "图1 智慧产业园区数字孪生平台总体业务流程图", width=Inches(5.8))
    add_body(doc, "系统功能模块如图2所示。图中红色边框表示本人负责分析、设计和实现的模块，包括安防管控、物业工单、企业服务、告警中心、系统管理以及登录鉴权等公共能力。其他模块由组员完成，本人负责接口约定、整体联调和验收。")
    add_figure(doc, ASSETS / "diagram-2-module-map.png", "图2 系统功能模块图（红色边框为本人负责模块）")
    add_body(doc, "安防管控模块用于集中展示摄像头、访客、车辆和门禁通行记录，并提供 1、4、9、16 分屏切换与异常通行、消防火警场景模拟。告警中心用于接收设备、能耗、安防、消防、环境和接口异常告警，支持确认、关闭、三维定位和生成关联工单。物业工单模块用于处理系统告警、企业报修和巡检异常，覆盖创建、派单、接单、处理、验收、完成和关闭状态。系统管理模块负责用户、角色、数据字典、接口、审计日志和虚拟仿真器设置。")

    add_heading(doc, "2.2 用例分析与设计", 2)
    add_body(doc, "本人负责模块的主要参与者包括系统管理员、安防值班人员、物业运维人员和企业服务人员。系统管理员维护用户、角色和日志；安防值班人员查看视频、确认告警并生成工单；物业运维人员接收和处理工单；企业服务人员维护企业档案与服务申请。组长负责模块用例关系如图3所示。")
    add_figure(doc, ASSETS / "diagram-3-use-case.png", "图3 组长负责模块用例图")
    add_body(doc, "告警生成工单是本系统最关键的业务用例，其需求拆解如表1所示。")
    add_table(
        doc,
        "表1 告警生成工单用例说明",
        ["项目", "说明"],
        [
            ["用例名称", "告警生成工单"],
            ["主要参与者", "安防值班人员、物业运维人员"],
            ["前置条件", "用户已登录并具有告警管理或工单管理权限；系统中存在未关闭告警"],
            ["基本流程", "值班人员查看告警详情，确认告警，选择生成工单；系统创建关联工单并记录告警编号、位置、优先级和时间轴；工单派发给物业运维人员处理"],
            ["异常流程", "告警已存在关联工单时，系统返回原工单，避免重复创建；权限不足时拒绝操作并返回提示"],
            ["后置条件", "告警状态进入处理中，工单进入待派单或待处理状态，操作写入审计日志"],
        ],
        [1.25, 4.85],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_body(doc, "权限模型采用 RBAC，即用户通过角色获得菜单权限和接口权限。RBAC 以角色作为用户与权限之间的稳定中间层，能够降低大型应用中逐用户授权的管理复杂度[2]。系统设置系统管理员、运营管理人员、物业运维人员、安防值班人员和能源管理人员等角色，并根据角色控制可见菜单和后端接口访问。")

    add_heading(doc, "2.3 实现技术", 2)
    add_body(doc, "系统采用前后端分离架构。前端使用 React 构建业务页面，Three.js WebGLRenderer 构建三维园区场景[5]；后端使用 Node.js 与 Express 提供 REST API；登录会话使用 JSON Web Token，JWT 是一种紧凑的、面向声明的安全信息传递格式[3]；实时数据使用 WebSocket 推送，WebSocket 协议支持客户端与服务器之间的双向通信[4]；数据持久化采用 JSON 文件，以保证虚拟演示环境能够零配置运行。系统技术架构如图4所示。")
    add_figure(doc, ASSETS / "diagram-4-architecture.png", "图4 系统技术架构图（红色边框为本人负责层）", width=Inches(5.8))
    add_body(doc, "具体技术选型如表2所示。")
    add_table(
        doc,
        "表2 系统技术选型",
        ["层级", "技术选型", "说明"],
        [
            ["前端", "React、Vite、Lucide React", "构建响应式管理后台、表单、表格和交互控件"],
            ["三维", "Three.js、WebGL", "构建园区楼栋、图层、标记点、日夜模式和对象拾取"],
            ["后端", "Node.js、Express", "提供 REST API、业务校验、权限控制和报表导出"],
            ["鉴权", "JWT、RBAC", "实现会话恢复、角色菜单权限和接口权限"],
            ["实时通道", "WebSocket", "推送虚拟设备遥测、告警和状态变化"],
            ["数据与测试", "JSON、Playwright、API Smoke Test", "零配置持久化，验证接口边界与关键用户流程"],
            ["版本管理", "Git、GitHub、Pull Request", "采用功能分支、阶段提交和 PR 汇总评审"],
        ],
        [0.85, 2.0, 3.25],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_body(doc, "在开发过程中使用 OpenAI Codex 作为 AI 协作工具，主要用于需求拆解、界面交互审计、测试脚本框架生成和文档整理。本人负责确认架构方向、检查生成结果、运行构建与测试、修复发现的问题并进行 Git 提交。AI 只承担辅助工作，最终设计决策、功能验收和版本管理由本人完成。")

    add_heading(doc, "2.4 功能模块设计", 2)
    add_body(doc, "告警与工单之间通过告警编号和工单编号建立关联。虚拟设备产生异常数据后，后端告警服务根据规则生成告警并通过 WebSocket 推送到前端；值班人员确认告警并生成工单；工单服务记录状态变化和处理时间轴；物业人员处理完成后，系统同步关闭关联告警。该闭环时序如图5所示。")
    add_figure(doc, ASSETS / "diagram-5-sequence.png", "图5 告警生成工单闭环时序图")
    add_body(doc, "工单状态机用于限制不合理的状态跳转，避免工单绕过处理步骤直接完成。状态转换规则如表3所示。系统管理员能够进行特殊维护操作，普通运维人员只能按照规定流程推进工单。")
    add_table(
        doc,
        "表3 工单状态转换规则",
        ["当前状态", "允许转换状态", "业务含义"],
        [
            ["待受理", "待派单、已驳回", "确认是否属于有效工单"],
            ["待派单", "待处理、已关闭", "分配责任班组或取消工单"],
            ["待处理", "处理中、已关闭", "处理人接单并开始处理"],
            ["处理中", "待验收、已完成", "提交处理结果或直接完成"],
            ["待验收", "已完成、处理中", "验收通过或退回整改"],
            ["已完成", "已关闭", "完成归档并关闭流程"],
        ],
        [1.0, 1.8, 3.3],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )

    chapter_three = add_heading(doc, "3 系统实现与测试", 1)
    chapter_three.paragraph_format.page_break_before = True
    add_heading(doc, "3.1 系统集成与首页驾驶舱", 2)
    add_body(doc, "作为组长，本人首先完成前后端公共数据约定和系统集成。后端提供 bootstrap 接口一次性返回楼栋、设备、告警、工单、企业、环境、权限和仿真器状态，前端在登录后加载数据并建立 WebSocket 连接。首页驾驶舱将企业入驻、空间出租、在线设备、当前告警、今日用电、待处理工单和环境评分集中展示，并支持跳转到对应业务模块，如图6所示。")
    add_figure(doc, ROOT / "verification-fullstack-dashboard.png", "图6 首页驾驶舱与业务集成界面")

    add_heading(doc, "3.2 登录鉴权与系统管理", 2)
    add_body(doc, "登录接口验证用户名和密码后签发 JWT，前端保存令牌并在后续请求中携带 Authorization 请求头。后端中间件解析令牌、恢复用户角色，并在接口层根据权限字符串判断是否允许访问。系统管理页面提供用户、角色、数据字典、接口、审计日志和仿真器设置，能够维护演示账号和角色权限，如图7所示。")
    add_figure(doc, ROOT / "verification-complete-system.png", "图7 系统管理与权限配置界面")
    add_body(doc, "系统内置五类演示账号：系统管理员、运营管理人员、物业运维人员、安防值班人员和能源管理人员。不同角色登录后可见菜单不同，例如物业运维人员只能访问驾驶舱、三维一张图、设备、环境、工单和告警模块，不能访问系统管理接口。")

    add_heading(doc, "3.3 安防管控实现", 2)
    add_body(doc, "安防管控模块将摄像头、访客、车辆、门禁和消防事件集中展示。视频区域支持 1、4、9、16 分屏切换，已有摄像头显示实时画面占位，未绑定摄像头的窗口明确显示为空闲窗口，避免重复伪造视频源。页面还提供异常通行和消防火警场景模拟，以便在没有真实安防设备的条件下演示事件处置流程，如图8所示。")
    add_figure(doc, ASSETS / "security-16-split.png", "图8 安防管控与视频 16 分屏界面")

    add_heading(doc, "3.4 告警中心与工单闭环实现", 2)
    add_body(doc, "告警中心统一展示设备、能耗、安防、消防、环境、门禁和接口异常告警，并根据等级显示不同状态。用户可以模拟实时告警、确认告警、三维定位、生成工单和关闭告警。生成工单时，系统自动继承告警来源、位置和等级，并按照严重程度设置工单优先级和 SLA。告警中心界面如图9所示。")
    add_figure(doc, ROOT / "verification-fullstack-alarms.png", "图9 统一告警中心与联动规则界面")
    add_body(doc, "物业工单页面使用看板和明细表展示待派单、待处理、处理中和待验收工单。用户可以创建工单、推进状态、退回验收、查看工单详情和流程时间轴。工单完成或关闭后，关联告警同步关闭，处理结果写入告警记录和审计日志，从而实现完整的可追溯闭环。")

    add_heading(doc, "3.5 版本管理与测试", 2)
    add_body(doc, "项目使用 Git 和 GitHub 进行版本管理。主分支 main 保存稳定基线，完整系统开发在 feature/full-virtual-demo-system 功能分支上进行，通过阶段性提交记录后端能力扩展、业务系统完成和交互质量收口。最终使用 Pull Request 汇总变更并检查是否可合并。主要提交包括 805486a“Expand virtual demo backend capabilities”、3bf8468“Complete virtual demo business system”和 63817de“Harden interactive demo workflows”。")
    add_body(doc, "测试分为构建测试、接口测试和端到端测试。构建测试验证前端能够正确打包；接口测试验证登录、数据加载、通用 CRUD、虚拟设备控制、场景仿真、权限边界、工单状态流转和报表导出；端到端测试使用 Playwright 验证登录、全局搜索、三维工具、视频分屏、工单详情和移动端布局。测试结果如表4所示。")
    add_table(
        doc,
        "表4 系统测试结果",
        ["编号", "测试内容", "预期结果", "实际结果"],
        [
            ["T01", "管理员登录与会话恢复", "正确进入首页并加载菜单", "通过"],
            ["T02", "物业运维人员访问系统管理接口", "返回 403，拒绝越权访问", "通过"],
            ["T03", "告警确认并生成关联工单", "告警进入处理中，创建工单", "通过"],
            ["T04", "普通角色非法跳转工单状态", "返回 400，阻止非法流转", "通过"],
            ["T05", "安防视频切换到 16 分屏", "页面显示 16 个稳定窗口", "通过"],
            ["T06", "Three.js 三维场景渲染", "Canvas 非空且可交互", "通过"],
            ["T07", "移动端页面布局", "无横向溢出，关键控件可操作", "通过"],
            ["T08", "报表 CSV 导出", "生成包含业务数据的文件", "通过"],
        ],
        [0.55, 2.2, 2.45, 0.9],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_body(doc, "实际验证中，npm run build、npm run test:api 和 npm run test:e2e 均通过。端到端测试检测到 Three.js Canvas 像素信号为 6144，移动端横向溢出为 0 像素。构建阶段仅存在 Three.js 相关包体积提示，不影响系统运行。")

    add_heading(doc, "4 总结", 1)
    add_body(doc, "本项目的主要难点不是单独实现一个展示页面，而是让告警、权限、工单和审计记录形成一致的业务状态。开发初期，部分按钮只有视觉效果，没有真实交互。本人通过逐项审计高频入口，补齐全局搜索、全屏控制、三维场景工具、视频分屏、工单详情和流程时间轴，并增加端到端测试，避免功能退化为“能看不能点”。")
    add_body(doc, "第二个难点是在没有本地设备和第三方平台的条件下保证系统可演示、可测试。解决方案是构建虚拟园区仿真器，提供日常运营、冷却泵故障、消防火警、能耗突增、地下泵房水浸和异常通行等场景。虚拟数据通过 WebSocket 推送，业务接口仍按照真实系统的方式进行权限校验、状态变更和持久化，因此后续接入真实 MQTT、视频平台或设备网关时可以保留主要业务流程。")
    add_body(doc, "第三个难点是多人协作中的集成与版本控制。本人通过功能分支、阶段提交、Pull Request 和自动化测试保持变更可追溯，并在组员模块完成后统一进行接口联调和页面验收。通过本项目，我进一步理解了领域软件工程中“先识别核心业务闭环，再选择技术实现”的重要性。")

    add_heading(doc, "5 项目团队与分工", 1)
    add_body(doc, "本项目由三名成员协作完成，按照组长 40%、成员二 30%、成员三 30% 的比例分配工作。第 1 节项目背景和第 2.1 节系统总体功能由小组共同讨论形成，各成员从个人负责模块开始独立分析、设计、实现和撰写。")
    add_table(
        doc,
        "表5 项目团队分工",
        ["成员", "比例", "主要负责内容"],
        [
            [f"组长：{PLACEHOLDER_NAME}", "40%", "总体架构、登录鉴权、安防管控、告警中心、物业工单、企业服务、系统管理、系统集成、自动化测试、GitHub 版本管理"],
            [f"成员二：{PLACEHOLDER_MEMBER_2}", "30%", "首页驾驶舱、数字孪生一张图、空间资产、Three.js 场景、图层控制、楼栋定位和响应式布局"],
            [f"成员三：{PLACEHOLDER_MEMBER_3}", "30%", "设备运维、能耗管理、环境监测、虚拟设备仿真、实时数据更新、报表导出"],
        ],
        [1.5, 0.65, 3.95],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT],
    )
    add_body(doc, f"组长 {PLACEHOLDER_NAME} 负责确定系统总体实现方向，设计并实现权限、安防、告警、工单和系统治理相关模块，组织前后端联调、测试和 GitHub 版本管理，并完成本个人报告。")
    add_body(doc, f"成员二 {PLACEHOLDER_MEMBER_2} 负责首页驾驶舱、数字孪生三维场景和空间资产模块，实现楼栋模型、图层控制、对象详情、搜索定位和移动端适配。")
    add_body(doc, f"成员三 {PLACEHOLDER_MEMBER_3} 负责设备运维、能耗管理、环境监测、虚拟仿真和报表功能，实现虚拟设备状态、能耗趋势、环境点位和数据导出。")

    doc.add_page_break()
    add_heading(doc, "6 附录：参考文献", 1)
    references = [
        "[1] 中国电信, 中国信息通信研究院. 5G空间计算白皮书[EB/OL]. 2023[2026-06-03]. https://shgis.com/post/1371.html.",
        "[2] Ferraiolo D, Kuhn R, Chandramouli R. Role-Based Access Control[M]. 2nd ed. Artech House, 2007.",
        "[3] Jones M, Bradley J, Sakimura N. JSON Web Token (JWT): RFC 7519[S]. IETF, 2015. https://www.rfc-editor.org/rfc/rfc7519.",
        "[4] Fette I, Melnikov A. The WebSocket Protocol: RFC 6455[S]. IETF, 2011. https://www.rfc-editor.org/rfc/rfc6455.",
        "[5] Three.js. WebGLRenderer Documentation[EB/OL]. [2026-06-03]. https://threejs.org/docs/pages/WebGLRenderer.html.",
        "[6] React. React Documentation[EB/OL]. [2026-06-03]. https://react.dev/.",
        "[7] Express.js. Express Documentation[EB/OL]. [2026-06-03]. https://expressjs.com/.",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(style="Bibliography" if "Bibliography" in [s.name for s in doc.styles] else None)
        paragraph.paragraph_format.first_line_indent = Pt(-24)
        paragraph.paragraph_format.left_indent = Pt(24)
        paragraph.paragraph_format.line_spacing = 1.25
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(reference)
        set_run_font(run, FONT_BODY, Pt(10.5))

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
